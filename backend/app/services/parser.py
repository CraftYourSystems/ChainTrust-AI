"""Document text extraction and normalisation.

Turns an uploaded PDF, DOCX, or TXT file into clean plain text suitable for
an LLM prompt. Knows nothing about FastAPI, HTTP, or LLMs — it takes bytes
and returns a string, so it can be exercised from a plain Python shell.

Public interface:

    extract_text(file, filename=None) -> str
    SUPPORTED_EXTENSIONS
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Callable

import docx
import pymupdf

from app.utils.exceptions import (
    DocumentParseError,
    EmptyDocumentError,
    InvalidFileTypeError,
)


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF via PyMuPDF.

    Image-only (scanned) PDFs have no text layer, so this returns an empty
    string for them. That is intentional: OCR is out of scope for the MVP,
    and extract_text() surfaces the empty result as EmptyDocumentError rather
    than attempting to recover it.
    """
    try:
        with pymupdf.open(stream=data, filetype="pdf") as document:
            if document.needs_pass:
                raise DocumentParseError("This PDF is password protected.")
            # Page break as paragraph break, so clauses don't run together.
            return "\n\n".join(page.get_text("text") for page in document)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("This PDF could not be opened. It may be corrupt.") from exc


def _extract_docx(data: bytes) -> str:
    """Extract paragraphs and table cells from a DOCX via python-docx."""
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError("This DOCX could not be opened. It may be corrupt.") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]

    # Contracts routinely put payment terms and SLAs in tables; skipping them
    # would drop commercially significant clauses.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    return "\n".join(blocks)


def _extract_txt(data: bytes) -> str:
    """Decode a plain-text file, tolerating the common Windows encodings."""
    # utf-8-sig also strips a BOM if one is present.
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise DocumentParseError("This text file uses an unrecognised character encoding.")


_EXTRACTORS: dict[str, Callable[[bytes], str]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
}

SUPPORTED_EXTENSIONS = frozenset(_EXTRACTORS)

_LINE_ENDINGS = re.compile(r"\r\n?")
_ZERO_WIDTH = re.compile("[\u200b-\u200d\ufeff]")
_INTRA_LINE_SPACE = re.compile(r"[^\S\n]+")
_REPEATED_BLANK_LINES = re.compile(r"\n{3,}")


def _clean(text: str) -> str:
    """Normalise whitespace without altering wording.

    Paragraph boundaries survive as a single blank line, and line starts are
    left untouched so clause numbering ("1.", "11.2", "(a)") is preserved.
    Nothing is summarised, rewritten, or truncated.
    """
    text = _LINE_ENDINGS.sub("\n", text)
    text = text.replace("\xa0", " ")
    text = _ZERO_WIDTH.sub("", text)

    lines = (_INTRA_LINE_SPACE.sub(" ", line).strip() for line in text.split("\n"))
    text = "\n".join(lines)

    return _REPEATED_BLANK_LINES.sub("\n\n", text).strip()


def _extension_of(name: str | None) -> str:
    """Return the validated lowercase extension for a filename."""
    if not name:
        raise InvalidFileTypeError("The file has no name, so its type cannot be determined.")

    extension = Path(name).suffix.lower()
    if extension not in _EXTRACTORS:
        described = extension or "unknown"
        raise InvalidFileTypeError(
            f"Unsupported file type '{described}'. Upload a PDF, DOCX, or TXT file."
        )
    return extension


def _resolve_source(file: Any, filename: str | None) -> tuple[bytes, str]:
    """Reduce any accepted input to (bytes, extension), reading the file once."""
    if isinstance(file, (str, Path)):
        path = Path(file)
        extension = _extension_of(filename or path.name)
        try:
            return path.read_bytes(), extension
        except OSError as exc:
            raise DocumentParseError(f"Could not read '{path.name}'.") from exc

    if isinstance(file, (bytes, bytearray, memoryview)):
        return bytes(file), _extension_of(filename)

    # Starlette's UploadFile, matched structurally rather than by import, so
    # this module stays free of any web-framework dependency.
    stream = getattr(file, "file", None)
    upload_name = getattr(file, "filename", None)
    if stream is not None and upload_name is not None:
        extension = _extension_of(filename or upload_name)
        stream.seek(0)
        return stream.read(), extension

    raise DocumentParseError("Unsupported input: pass bytes, a path, or an uploaded file.")


import logging

logger = logging.getLogger("chaintrust.parser")


def extract_text(file: Any, filename: str | None = None) -> str:
    """Extract clean, LLM-ready text from a contract document.

    Args:
        file: An UploadFile, raw bytes, or a filesystem path.
        filename: Overrides the name used for type detection. Required when
            ``file`` is raw bytes.

    Returns:
        Normalised UTF-8 text, never empty.

    Raises:
        InvalidFileTypeError: Extension is not PDF, DOCX, or TXT.
        DocumentParseError: File is corrupt, encrypted, or unreadable.
        EmptyDocumentError: File holds no extractable text.
    """
    data, extension = _resolve_source(file, filename)

    if not data:
        raise EmptyDocumentError("The uploaded file is empty.")

    text = _clean(_EXTRACTORS[extension](data))

    if not text:
        raise EmptyDocumentError()

    char_count = len(text)
    first_500 = text[:500]
    last_500 = text[-500:] if char_count > 500 else text

    logger.info(
        "\n========================================\n"
        "--- CONTRACT TEXT EXTRACTION AUDIT ---\n"
        "  - File Name: %s\n"
        "  - Extracted Character Count: %d\n"
        "  - First 500 chars:\n%s\n\n"
        "  - Last 500 chars:\n%s\n"
        "========================================",
        filename or "unknown",
        char_count,
        first_500,
        last_500,
    )

    return text
